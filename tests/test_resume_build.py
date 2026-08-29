"""Tests for resume_build.py's DOCX construction. Uses python-docx's own reader to verify
output, not mocks -- building a real (tiny) DOCX in a temp dir is fast and exercises the
real library integration, matching how gmail.py's tests exercise real MIME construction."""

import os
from unittest.mock import MagicMock

import pytest
from docx import Document

import config
import resume_build

_STRATEGY = {
    "section_order": ["Experience", "Projects", "Skills", "Leadership"],
    "projects_included": ["Viant", "Hiya"],
    "skills_groups": [{"label": "Data & Tools", "skills": ["SQL", "Python"]}],
}
_MASTER = {
    "name": "Kishore Theeraj Vasudevan Jaya",
    "contact": {"location": "Hanover, NH", "phone": "+1 603-322-0535",
                "email": "kishore@example.com", "linkedin": "linkedin.com/in/kishoretheeraj"},
    "roles": [
        {"title": "Associate Product Manager", "company": "Protium Finance",
         "descriptor": "Digital lending company", "location": "Bangalore, India",
         "period": "Apr 2025 - Aug 2025", "bullets": ["Eliminated vendor cost via a build-vs-buy model."]},
    ],
    "education": [
        {"institution": "Dartmouth College", "location": "Hanover, NH",
         "program": "Master's, Engineering Management", "graduation": "Nov 2026", "coursework": None},
    ],
    "projects": {
        "Viant": {"descriptor": None, "period": None, "bullets": ["Narrowed 7 vendors to 2 finalists."]},
        "Hiya": {"descriptor": "Dartmouth Project", "period": "2025", "bullets": ["Protected 500M+ users."]},
    },
    "leadership_bullets": ["Mentored 500+ students."],
}


def test_build_docx_creates_file_at_output_path(tmp_path):
    output_path = str(tmp_path / "resume.docx")
    result = resume_build.build_docx(_STRATEGY, _MASTER, output_path)
    assert result == output_path
    assert os.path.exists(output_path)


def test_build_docx_includes_role_and_project_content():
    doc_path = "test_build_docx_includes_role_and_project_content.docx"
    try:
        resume_build.build_docx(_STRATEGY, _MASTER, doc_path)
        doc = Document(doc_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "KISHORE THEERAJ VASUDEVAN JAYA" in full_text
        assert "Hanover, NH" in full_text and "kishore@example.com" in full_text
        assert "Associate Product Manager" in full_text
        assert "Protium Finance" in full_text
        assert "Bangalore, India" in full_text
        assert "Viant" in full_text
        assert "Narrowed 7 vendors to 2 finalists." in full_text
        assert "Data & Tools" in full_text and "SQL" in full_text
        assert "Mentored 500+ students." in full_text
    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)


def test_build_docx_uses_right_tab_stop_for_dates_and_locations():
    doc_path = "test_build_docx_uses_right_tab_stop_for_dates_and_locations.docx"
    try:
        resume_build.build_docx(_STRATEGY, _MASTER, doc_path)
        doc = Document(doc_path)
        company_line = next(p for p in doc.paragraphs if "Protium Finance" in p.text)
        tabs = list(company_line.paragraph_format.tab_stops)
        assert len(tabs) == 1
        from docx.enum.text import WD_TAB_ALIGNMENT
        assert tabs[0].alignment == WD_TAB_ALIGNMENT.RIGHT
        assert "\t" in company_line.text
    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)


def test_build_docx_section_headers_have_bottom_border():
    from docx.oxml.ns import qn
    doc_path = "test_build_docx_section_headers_have_bottom_border.docx"
    try:
        resume_build.build_docx(_STRATEGY, _MASTER, doc_path)
        doc = Document(doc_path)
        header = next(p for p in doc.paragraphs if p.text.strip() == "EXPERIENCE")
        pPr = header._p.pPr
        assert pPr.find(qn("w:pBdr")) is not None
    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)


def test_build_docx_raises_on_unknown_section_name():
    bad_strategy = dict(_STRATEGY, section_order=["Summary"])
    with pytest.raises(ValueError, match="unknown section"):
        resume_build.build_docx(bad_strategy, _MASTER, "should_not_be_created.docx")
    assert not os.path.exists("should_not_be_created.docx")


def test_build_docx_applies_named_margin_preset():
    from docx.shared import Twips
    doc_path = "test_build_docx_applies_named_margin_preset.docx"
    try:
        resume_build.build_docx(_STRATEGY, _MASTER, doc_path, margin_preset="tight")
        doc = Document(doc_path)
        section = doc.sections[0]
        preset = resume_build._MARGIN_PRESETS["tight"]
        assert section.top_margin == Twips(preset["top"])
        assert section.left_margin == Twips(preset["left"])
    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)


def test_margin_ladder_covers_every_preset():
    assert set(resume_build._MARGIN_LADDER) == set(resume_build._MARGIN_PRESETS.keys())


def test_standard_preset_is_in_margin_ladder():
    assert "standard" in resume_build._MARGIN_LADDER


# ── convert_to_pdf ─────────────────────────────────────────────────────────────

def test_convert_to_pdf_calls_soffice_and_returns_pdf_path(mocker, tmp_path):
    run = mocker.patch("resume_build.subprocess.run")
    docx_path = str(tmp_path / "resume.docx")
    result = resume_build.convert_to_pdf(docx_path, str(tmp_path))
    run.assert_called_once()
    args = run.call_args[0][0]
    assert args[0] == "soffice"
    assert "--headless" in args
    assert docx_path in args
    assert result == str(tmp_path / "resume.pdf")


def test_convert_to_pdf_uses_configured_timeout(mocker, tmp_path):
    run = mocker.patch("resume_build.subprocess.run")
    resume_build.convert_to_pdf(str(tmp_path / "r.docx"), str(tmp_path))
    assert run.call_args.kwargs["timeout"] == config.RESUME_SOFFICE_TIMEOUT_SECONDS


# ── page_count ─────────────────────────────────────────────────────────────────

def test_page_count_reads_pdf_page_count(mocker):
    fake_reader = MagicMock()
    fake_reader.pages = [MagicMock(), MagicMock()]
    mocker.patch("resume_build.PdfReader", return_value=fake_reader)
    assert resume_build.page_count("fake.pdf") == 2


# ── fit_to_one_page ────────────────────────────────────────────────────────────

def test_fit_to_one_page_returns_immediately_when_already_one_page(mocker, tmp_path):
    mocker.patch.object(resume_build, "build_docx", return_value="r.docx")
    mocker.patch.object(resume_build, "convert_to_pdf", return_value="r.pdf")
    mocker.patch.object(resume_build, "page_count", return_value=1)
    pdf_path, preset = resume_build.fit_to_one_page(_STRATEGY, _MASTER, "r.docx", str(tmp_path))
    assert pdf_path == "r.pdf"
    assert preset == "standard"


def test_fit_to_one_page_walks_margin_ladder_until_it_fits(mocker, tmp_path):
    mocker.patch.object(resume_build, "build_docx", return_value="r.docx")
    mocker.patch.object(resume_build, "convert_to_pdf", return_value="r.pdf")
    # 5 rungs total (line/bullet/header spacing don't change page count in this fake sequence,
    # margins rung on attempt 4 finally fits)
    mocker.patch.object(resume_build, "page_count", side_effect=[2, 2, 2, 1])
    pdf_path, preset = resume_build.fit_to_one_page(_STRATEGY, _MASTER, "r.docx", str(tmp_path))
    assert pdf_path == "r.pdf"
    assert preset in resume_build._MARGIN_LADDER


def test_fit_to_one_page_raises_still_overflow_error_after_exhausting_ladder(mocker, tmp_path):
    mocker.patch.object(resume_build, "build_docx", return_value="r.docx")
    mocker.patch.object(resume_build, "convert_to_pdf", return_value="r.pdf")
    mocker.patch.object(resume_build, "page_count", return_value=2)
    with pytest.raises(resume_build.StillOverflowError):
        resume_build.fit_to_one_page(_STRATEGY, _MASTER, "r.docx", str(tmp_path))


def test_fit_rungs_actually_tighten_progressively():
    # Regression test: an earlier version named 5 rungs but only "margins" changed anything, and
    # even that advanced only one step (standard -> tight), never reaching aggressive/floor --
    # found when real content that needed the full ladder still overflowed one page. Each rung
    # must be strictly no looser than the previous one on every dimension that varies.
    for prev, cur in zip(resume_build._FIT_RUNGS, resume_build._FIT_RUNGS[1:]):
        assert sum(cur["header_spacing"]) <= sum(prev["header_spacing"])
        assert sum(cur["bullet_spacing"]) <= sum(prev["bullet_spacing"])
        assert cur["entry_spacing_before"] <= prev["entry_spacing_before"]
        assert cur["body_size"] <= prev["body_size"]
    assert resume_build._FIT_RUNGS[-1]["margin_preset"] == "floor"


def test_build_docx_groups_consecutive_same_company_roles_under_one_header(tmp_path):
    # Regression test: the real corpus template prints a company's name/descriptor/location line
    # once for a multi-role tenure (e.g. 4 promotions at the same employer), not once per role --
    # repeating it made a live-rendered resume visually read as 4 separate jobs, not one tenure.
    master = dict(_MASTER, roles=[
        {"title": "Intern", "company": "Acme", "descriptor": "Fintech", "location": "Bangalore, India",
         "period": "2022", "bullets": ["Did intern things."]},
        {"title": "Analyst", "company": "Acme", "descriptor": "Fintech", "location": "Bangalore, India",
         "period": "2023", "bullets": ["Did analyst things."]},
    ])
    strategy = dict(_STRATEGY, section_order=["Experience"])
    output_path = str(tmp_path / "grouped.docx")
    resume_build.build_docx(strategy, master, output_path)
    doc = Document(output_path)
    company_lines = [p for p in doc.paragraphs if "Acme" in p.text]
    assert len(company_lines) == 1
    assert any(p.text.strip() == "Analyst\t2023" for p in doc.paragraphs)


def test_build_docx_applies_style_tightening_to_bullets_and_headers(tmp_path):
    output_path = str(tmp_path / "tight.docx")
    resume_build.build_docx(_STRATEGY, _MASTER, output_path, style=resume_build._FIT_RUNGS[-1])
    doc = Document(output_path)
    bullet = next(p for p in doc.paragraphs if "Eliminated vendor cost" in p.text)
    assert bullet.paragraph_format.space_before.pt == 0
    assert bullet.runs[0].font.size.pt == 9
