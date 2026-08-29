"""
Builds a resume DOCX from a strategy (resume_agent.py's stage-4 output) and the
master resume data (resume/data/master.json). Pure/deterministic -- no LLM calls,
no I/O beyond writing the output file. Raises on failure; does not swallow
exceptions, since this is a manual, interactive tool. See
docs/superpowers/specs/2026-08-29-phase3-resume-intelligence-design.md.
"""

import os
import subprocess

from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pypdf import PdfReader

import config

# ── Typography and spacing presets ──────────────────────────────────────────────
# Margin values extracted from the user's own real resume corpus (77 .docx files, checked
# 2026-08-29): "standard" is the single most common historical combo (13/77 files) and matches
# the ATS-safe floor (0.5in) from current best-practice guidance. "floor" is the exact margin
# combo used by the most recently refined real template when a resume needed to be squeezed to
# one page (asymmetric on purpose -- top/bottom gets squeezed before left/right, since page height
# is what's actually overflowing).

_MARGIN_PRESETS = {
    "comfortable": {"top": 900, "bottom": 900, "left": 900, "right": 900},
    "standard": {"top": 720, "bottom": 720, "left": 720, "right": 720},
    "tight": {"top": 576, "bottom": 576, "left": 576, "right": 576},
    "aggressive": {"top": 504, "bottom": 432, "left": 576, "right": 576},
    "floor": {"top": 432, "bottom": 360, "left": 576, "right": 576},
}
_MARGIN_LADDER = ["comfortable", "standard", "tight", "aggressive", "floor"]

_NAME_SIZE_PT = 13
_SECTION_HEADER_SIZE_PT = 10
_BODY_SIZE_PT = 10
_FONT_NAME = config.RESUME_FONT_NAME

# ── Fitting-ladder style presets ────────────────────────────────────────────────
# Each rung cumulatively tightens spacing/font beyond the previous one. "margin_preset" is a
# _MARGIN_PRESETS key; header_spacing/bullet_spacing/entry_spacing_before are point values;
# body_size is the point size for bullets, entry lines, skills, and leadership text (section
# headers and the name stay fixed regardless of rung -- only body-level text shrinks).
_DEFAULT_STYLE = {
    "margin_preset": "standard", "header_spacing": (6, 2), "bullet_spacing": (1, 1),
    "entry_spacing_before": 4, "body_size": _BODY_SIZE_PT,
}
_FIT_RUNGS = [
    {"margin_preset": "standard", "header_spacing": (6, 2), "bullet_spacing": (1, 1), "entry_spacing_before": 4, "body_size": 10},
    {"margin_preset": "standard", "header_spacing": (4, 1), "bullet_spacing": (0, 1), "entry_spacing_before": 3, "body_size": 10},
    {"margin_preset": "tight", "header_spacing": (3, 1), "bullet_spacing": (0, 0), "entry_spacing_before": 2, "body_size": 10},
    {"margin_preset": "aggressive", "header_spacing": (2, 1), "bullet_spacing": (0, 0), "entry_spacing_before": 1, "body_size": 9},
    {"margin_preset": "floor", "header_spacing": (2, 1), "bullet_spacing": (0, 0), "entry_spacing_before": 0, "body_size": 9},
]


# ── Low-level formatting helpers ────────────────────────────────────────────────

def _apply_margins(doc, preset_name):
    preset = _MARGIN_PRESETS[preset_name]
    section = doc.sections[0]
    section.top_margin = Twips(preset["top"])
    section.bottom_margin = Twips(preset["bottom"])
    section.left_margin = Twips(preset["left"])
    section.right_margin = Twips(preset["right"])


def _content_width_twips(doc):
    section = doc.sections[0]
    return section.page_width.twips - section.left_margin.twips - section.right_margin.twips


def _add_bottom_border(paragraph):
    """Add a single bottom-border rule under a paragraph -- the real corpus's own section-header
    technique (present in every file checked, both the ALL-CAPS/Calibri majority and the
    Title-Case/Garamond variant). This is what makes a header read as a header, not just bold
    text sitting in the body."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_right_tab_stop(paragraph, position_twips):
    paragraph.paragraph_format.tab_stops.add_tab_stop(Twips(position_twips), WD_TAB_ALIGNMENT.RIGHT)


def _new_paragraph(doc, style=None):
    """python-docx's default template applies 1.15x line spacing and 10pt space_after to every
    paragraph unless overridden (found live: this alone was enough extra height, spread across
    ~20 paragraphs, to push a one-page resume onto a second page even at the tightest fitting-
    ladder rung). Every paragraph in this module starts from a zeroed, single-spaced baseline;
    callers then set whatever space_before/space_after the real corpus pattern calls for."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def _run(paragraph, text, bold=False, italic=False, size=_BODY_SIZE_PT):
    run = paragraph.add_run(text)
    run.font.name = _FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


# ── Header (name + contact line) ────────────────────────────────────────────────

def _add_name_heading(doc, name):
    p = _new_paragraph(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, name.upper(), bold=True, size=_NAME_SIZE_PT)
    p.paragraph_format.space_after = Pt(1)


def _add_contact_line(doc, contact):
    parts = [v for v in (
        contact.get("location"), contact.get("phone"), contact.get("email"), contact.get("linkedin"),
    ) if v]
    if not parts:
        return
    p = _new_paragraph(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "  |  ".join(parts))
    p.paragraph_format.space_after = Pt(6)


# ── Section header (bordered, corpus-wide pattern) ──────────────────────────────

def _add_section_header(doc, title, style):
    p = _new_paragraph(doc)
    _run(p, title.upper(), bold=True, size=_SECTION_HEADER_SIZE_PT)
    before, after = style["header_spacing"]
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    _add_bottom_border(p)


# ── Entry lines (two-line pattern: roles/education; one-line pattern: projects) ─

def _add_two_line_entry(doc, content_width, primary, descriptor, location, secondary, period, style):
    """Roles and education both render as: bold primary [+ plain descriptor], tab, location on
    line 1; bold+italic secondary [title/program], tab, italic period on line 2. Matches the
    right-tab-stopped pattern found in every file checked in the user's real resume corpus."""
    size = style["body_size"]
    p1 = _new_paragraph(doc)
    _run(p1, primary, bold=True, size=size)
    if descriptor:
        _run(p1, f" | {descriptor}", size=size)
    if location:
        _add_right_tab_stop(p1, content_width)
        _run(p1, f"\t{location}", size=size)
    p1.paragraph_format.space_before = Pt(style["entry_spacing_before"])

    p2 = _new_paragraph(doc)
    _run(p2, secondary, bold=True, italic=True, size=size)
    if period:
        _add_right_tab_stop(p2, content_width)
        _run(p2, f"\t{period}", italic=True, size=size)


def _add_one_line_entry(doc, content_width, primary, descriptor, period, style):
    """Projects render as a single line: bold name [+ plain descriptor], tab, period on the right
    -- no separate title/date line the way roles and education have."""
    size = style["body_size"]
    p = _new_paragraph(doc)
    _run(p, primary, bold=True, size=size)
    if descriptor:
        _run(p, f" | {descriptor}", size=size)
    if period:
        _add_right_tab_stop(p, content_width)
        _run(p, f"\t{period}", size=size)
    p.paragraph_format.space_before = Pt(style["entry_spacing_before"])


def _add_bullet(doc, text, style):
    size = style["body_size"]
    before, after = style["bullet_spacing"]
    p = _new_paragraph(doc, style="List Bullet")
    _run(p, text, size=size)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


# ── Section builders ───────────────────────────────────────────────────────────

def _add_role_title_line(doc, content_width, title, period, style):
    """A role's title/date line only -- used for the 2nd+ consecutive role at the same company,
    where the company/descriptor/location header (see _add_experience_section) isn't repeated."""
    size = style["body_size"]
    p = _new_paragraph(doc)
    _run(p, title, bold=True, italic=True, size=size)
    if period:
        _add_right_tab_stop(p, content_width)
        _run(p, f"\t{period}", italic=True, size=size)
    p.paragraph_format.space_before = Pt(style["entry_spacing_before"])


def _add_experience_section(doc, master, content_width, style):
    """Consecutive roles at the same company (matched on company+descriptor+location) share one
    company header line, matching the real corpus's own pattern -- printing it once per role
    when someone was promoted internally multiple times reads as separate jobs, not one tenure."""
    _add_section_header(doc, "Experience", style)
    last_company_key = None
    for role in master.get("roles", []):
        company_key = (role["company"], role.get("descriptor"), role.get("location"))
        if company_key != last_company_key:
            _add_two_line_entry(
                doc, content_width, role["company"], role.get("descriptor"), role.get("location"),
                role["title"], role["period"], style,
            )
            last_company_key = company_key
        else:
            _add_role_title_line(doc, content_width, role["title"], role["period"], style)
        for bullet in role.get("bullets", []):
            _add_bullet(doc, bullet, style)


def _add_education_section(doc, master, content_width, style):
    education = master.get("education")
    if not education:
        return
    _add_section_header(doc, "Education", style)
    for entry in education:
        _add_two_line_entry(
            doc, content_width, entry["institution"], None, entry.get("location"),
            entry["program"], entry["graduation"], style,
        )
        if entry.get("coursework"):
            p = _new_paragraph(doc)
            _run(p, entry["coursework"], size=style["body_size"])


def _add_projects_section(doc, master, content_width, projects_included, style):
    _add_section_header(doc, "Projects", style)
    for project_name in projects_included:
        project = master.get("projects", {}).get(project_name)
        if not project:
            continue
        _add_one_line_entry(
            doc, content_width, project_name, project.get("descriptor"), project.get("period"), style,
        )
        for bullet in project.get("bullets", []):
            _add_bullet(doc, bullet, style)


def _add_skills_section(doc, master, skills_groups, style):
    _add_section_header(doc, "Skills", style)
    for group in skills_groups:
        p = _new_paragraph(doc)
        _run(p, f"{group['label']}: ", bold=True, size=style["body_size"])
        _run(p, ", ".join(group["skills"]), size=style["body_size"])


def _add_leadership_section(doc, master, style):
    _add_section_header(doc, "Leadership", style)
    for bullet in master.get("leadership_bullets", []):
        _add_bullet(doc, bullet, style)


_SECTION_BUILDERS = {
    "Experience": lambda doc, master, content_width, strategy, style: _add_experience_section(
        doc, master, content_width, style,
    ),
    "Education": lambda doc, master, content_width, strategy, style: _add_education_section(
        doc, master, content_width, style,
    ),
    "Projects": lambda doc, master, content_width, strategy, style: _add_projects_section(
        doc, master, content_width, strategy.get("projects_included", []), style,
    ),
    "Skills": lambda doc, master, content_width, strategy, style: _add_skills_section(
        doc, master, strategy.get("skills_groups", []), style,
    ),
    "Leadership": lambda doc, master, content_width, strategy, style: _add_leadership_section(
        doc, master, style,
    ),
}


def build_docx(strategy, master, output_path, margin_preset="standard", style=None):
    """Build a resume DOCX from `strategy` (section order, projects included, skills groups) and
    `master` (resolved resume/data/master.json content -- see resume_agent._resolve_master).
    `style` (header_spacing/bullet_spacing/body_size) defaults to _DEFAULT_STYLE; fit_to_one_page
    passes progressively tighter values from _FIT_RUNGS as it walks the ladder. Returns
    output_path. Raises ValueError if strategy["section_order"] names a section outside
    config.RESUME_ALLOWED_SECTIONS -- an unconstrained value here previously meant a whole section
    silently vanished from the built resume instead of failing loudly."""
    style = style or _DEFAULT_STYLE
    doc = Document()
    _apply_margins(doc, margin_preset)
    content_width = _content_width_twips(doc)

    _add_name_heading(doc, master.get("name", ""))
    _add_contact_line(doc, master.get("contact", {}))

    for section_name in strategy.get("section_order", []):
        if section_name not in config.RESUME_ALLOWED_SECTIONS:
            raise ValueError(
                f"strategy proposed unknown section {section_name!r} -- must be one of "
                f"{config.RESUME_ALLOWED_SECTIONS}"
            )
        _SECTION_BUILDERS[section_name](doc, master, content_width, strategy, style)

    doc.save(output_path)
    return output_path


# ── PDF conversion ──────────────────────────────────────────────────────────────

class StillOverflowError(Exception):
    """Raised when the deterministic fitting-ladder rungs can't get a resume to one page.
    Caller (resume_agent.py) catches this and triggers a content-editing regeneration."""


def convert_to_pdf(docx_path, output_dir):
    """Convert docx_path to PDF via LibreOffice headless. Returns the output PDF path. Raises on
    any failure (missing soffice binary, conversion error, timeout) -- never swallowed."""
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
        check=True, capture_output=True, timeout=config.RESUME_SOFFICE_TIMEOUT_SECONDS,
    )
    base = os.path.splitext(os.path.basename(docx_path))[0]
    return os.path.join(output_dir, base + ".pdf")


def page_count(pdf_path):
    return len(PdfReader(pdf_path).pages)


# ── Fitting ladder (corpus spec Part 13, formatting rungs only -- see Task 9's docstring note) ──
# _FIT_RUNGS (defined above, alongside _DEFAULT_STYLE) is the real, cumulative implementation:
# each successive rung tightens header/bullet spacing, margins, and (at the last two rungs) body
# font size, matching the named progression (line spacing -> bullet spacing -> header spacing ->
# margins -> font floor). An earlier version named these rungs but only margins actually changed
# anything, and even that advanced just one step -- found when real, denser strategy output
# (4 roles + 3 projects + a skills section) still overflowed one page at the loosest tightening.


def fit_to_one_page(strategy, master, output_path, output_dir):
    """
    Build, convert, and check page count, walking _FIT_RUNGS in order until the PDF is one page.
    Returns (pdf_path, margin_preset_used). Raises StillOverflowError if still >1 page after
    every rung -- the caller should treat that as a signal to shorten content, not retry
    formatting again.
    """
    for rung in _FIT_RUNGS:
        docx_path = build_docx(strategy, master, output_path, margin_preset=rung["margin_preset"], style=rung)
        pdf_path = convert_to_pdf(docx_path, output_dir)
        if page_count(pdf_path) <= 1:
            return pdf_path, rung["margin_preset"]
    raise StillOverflowError(
        f"still overflows one page after every formatting rung (final preset: "
        f"{_FIT_RUNGS[-1]['margin_preset']})"
    )
